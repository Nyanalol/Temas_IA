"""
cv_writer.py — Genera un archivo .docx a partir de un modelo CVDocument.

Usa python-docx para producir un Word bien formateado listo para enviar.
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from .models import CVDocument

_SECTION_LABELS = {
    "profile": {"es": "PERFIL PROFESIONAL", "en": "PROFESSIONAL PROFILE"},
    "experience": {"es": "EXPERIENCIA PROFESIONAL", "en": "PROFESSIONAL EXPERIENCE"},
    "education": {"es": "EDUCACIÓN", "en": "EDUCATION"},
    "skills": {"es": "HABILIDADES TÉCNICAS", "en": "TECHNICAL SKILLS"},
    "languages": {"es": "IDIOMAS", "en": "LANGUAGES"},
    "certifications": {"es": "CERTIFICACIONES", "en": "CERTIFICATIONS"},
}


def _label(section: str, lang: str) -> str:
    return _SECTION_LABELS.get(section, {}).get(lang, section.upper())


def write_cv_docx(cv: CVDocument, output_path: Path) -> Path:
    """
    Genera un .docx a partir de un CVDocument y lo guarda en output_path.
    Devuelve la ruta del archivo generado.
    """
    lang = cv.language if cv.language in ("es", "en") else "es"
    doc = Document()

    # ── Nombre ────────────────────────────────────────────────────────────────
    title = doc.add_heading(cv.contact.name, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Contacto ──────────────────────────────────────────────────────────────
    contact_parts = [
        p
        for p in [
            cv.contact.email,
            cv.contact.phone,
            cv.contact.location,
            cv.contact.linkedin,
            cv.contact.github,
        ]
        if p
    ]
    if contact_parts:
        para = doc.add_paragraph(" · ".join(contact_parts))
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.runs[0].font.size = Pt(10)

    doc.add_paragraph()

    # ── Resumen ───────────────────────────────────────────────────────────────
    if cv.summary:
        doc.add_heading(_label("profile", lang), level=1)
        doc.add_paragraph(cv.summary)
        doc.add_paragraph()

    # ── Experiencia ───────────────────────────────────────────────────────────
    if cv.experience:
        doc.add_heading(_label("experience", lang), level=1)
        for exp in cv.experience:
            p = doc.add_paragraph()
            run = p.add_run(f"{exp.role} — {exp.company}")
            run.bold = True
            period_line = exp.period
            if exp.location:
                period_line += f" · {exp.location}"
            doc.add_paragraph(period_line).runs[0].font.size = Pt(10)
            for achievement in exp.achievements:
                doc.add_paragraph(f"• {achievement}", style="List Bullet")
            doc.add_paragraph()

    # ── Educación ─────────────────────────────────────────────────────────────
    if cv.education:
        doc.add_heading(_label("education", lang), level=1)
        for edu in cv.education:
            p = doc.add_paragraph()
            run = p.add_run(f"{edu.degree} — {edu.institution}")
            run.bold = True
            doc.add_paragraph(edu.period).runs[0].font.size = Pt(10)
            if edu.details:
                doc.add_paragraph(edu.details)
            doc.add_paragraph()

    # ── Skills ────────────────────────────────────────────────────────────────
    if cv.skills:
        doc.add_heading(_label("skills", lang), level=1)
        doc.add_paragraph(" · ".join(cv.skills))
        doc.add_paragraph()

    # ── Idiomas ───────────────────────────────────────────────────────────────
    if cv.languages:
        doc.add_heading(_label("languages", lang), level=1)
        for language in cv.languages:
            doc.add_paragraph(f"• {language}", style="List Bullet")
        doc.add_paragraph()

    # ── Certificaciones ───────────────────────────────────────────────────────
    if cv.certifications:
        doc.add_heading(_label("certifications", lang), level=1)
        for cert in cv.certifications:
            doc.add_paragraph(f"• {cert}", style="List Bullet")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"[cv_writer] CV guardado en: {output_path}")
    return output_path
